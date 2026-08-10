from pathlib import Path
import subprocess
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / ".codex-temp" / "documentation"
OUTPUT = ROOT / "docs" / "PlyLegal_Client_Portal_Technical_Documentation.docx"
ARCHITECTURE_PNG = WORK / "architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F6F9"
GRAY = "5F6B76"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD5E1"
RED = "9B1C1C"
PALE_RED = "FDECEC"
GREEN = "1D6B62"
PALE_GREEN = "E9F5F2"
GOLD = "7A5A00"
PALE_GOLD = "FFF8E5"
BLACK = "1F2937"
WHITE = "FFFFFF"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(value):
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def apply_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        set_row_cant_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    next_abstract = max(abstract_ids or [0]) + 1
    next_num = max(num_ids or [0]) + 1

    def create(kind, abstract_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create("bullet", next_abstract, next_num)
    create("decimal", next_abstract + 1, next_num + 1)
    return next_num, next_num + 1


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(DARK_BLUE)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, MID_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Code Block" not in [style.name for style in styles]:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = rgb(BLACK)
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.05

    if "Small Note" not in [style.name for style in styles]:
        small = styles.add_style("Small Note", 1)
    else:
        small = styles["Small Note"]
    small.font.name = "Calibri"
    small._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    small._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    small.font.size = Pt(9)
    small.font.color.rgb = rgb(GRAY)
    small.paragraph_format.space_before = Pt(2)
    small.paragraph_format.space_after = Pt(6)
    small.paragraph_format.line_spacing = 1.15


def set_paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill, accent)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=BLACK)
    return p


def add_bullet(doc, bullet_num_id, text, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        set_run_font(p.add_run(text))
    return p


def add_numbered(doc, number_num_id, text):
    p = doc.add_paragraph()
    apply_num(p, number_num_id)
    set_run_font(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths_dxa, font_size=8.8, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr()
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(str(header))
        set_run_font(run, size=font_size, color=DARK_BLUE, bold=True)

    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run("" if value is None else str(value))
            set_run_font(run, size=font_size, color=BLACK, bold=(first_col_bold and idx == 0))
    apply_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, LIGHT_GRAY, BORDER)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=8.5, color=BLACK)
    return p


def add_section_break(doc):
    doc.add_page_break()


def set_picture_alt(inline_shape, description):
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)


def make_architecture_diagram(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 850), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    regular = ImageFont.truetype(str(font_path), 30)
    small = ImageFont.truetype(str(font_path), 22)
    bold = ImageFont.truetype(str(bold_path), 34)
    title = ImageFont.truetype(str(bold_path), 42)

    def box(xy, fill, outline, heading, lines):
        draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=4)
        x1, y1, x2, y2 = xy
        draw.text(((x1 + x2) / 2, y1 + 34), heading, font=bold, fill="#0B2545", anchor="ma")
        y = y1 + 92
        for line in lines:
            draw.text(((x1 + x2) / 2, y), line, font=small, fill="#263645", anchor="ma")
            y += 34

    def arrow(start, end, label=None):
        draw.line([start, end], fill="#4F6B7A", width=6)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) > abs(ey - sy):
            direction = 1 if ex > sx else -1
            points = [(ex, ey), (ex - 18 * direction, ey - 12), (ex - 18 * direction, ey + 12)]
        else:
            direction = 1 if ey > sy else -1
            points = [(ex, ey), (ex - 12, ey - 18 * direction), (ex + 12, ey - 18 * direction)]
        draw.polygon(points, fill="#4F6B7A")
        if label:
            mx, my = (sx + ex) / 2, (sy + ey) / 2
            bbox = draw.textbbox((0, 0), label, font=small)
            pad = 8
            draw.rounded_rectangle(
                (mx - (bbox[2] - bbox[0]) / 2 - pad, my - 20, mx + (bbox[2] - bbox[0]) / 2 + pad, my + 20),
                radius=8,
                fill="white",
            )
            draw.text((mx, my), label, font=small, fill="#4F6B7A", anchor="mm")

    draw.text((800, 45), "Current application architecture", font=title, fill="#0B2545", anchor="ma")
    box((60, 150, 430, 385), "#F4F6F9", "#2E74B5", "Users", [
        "Client browser",
        "Admin portal",
        "Zoho CRM widget",
    ])
    box((615, 125, 985, 410), "#E8EEF5", "#2E74B5", "Next.js portal", [
        "App Router pages",
        "Route handlers (/api)",
        "Valtio + React Query",
        "Validation and navigation",
    ])
    box((1100, 150, 1540, 385), "#E9F5F2", "#1D6B62", "Zoho CRM", [
        "Contacts and Deals",
        "Dependants / documents",
        "Corrections and references",
        "Uploaded attachments",
    ])
    box((260, 550, 700, 790), "#FFF8E5", "#7A5A00", "Firebase", [
        "Authentication",
        "Cloud Firestore",
        "Profiles, drafts, progress",
        "Conversations and resources",
    ])
    box((900, 550, 1340, 790), "#FDECEC", "#9B1C1C", "External helpers", [
        "Zoho access-token broker",
        "Password-reset email webhook",
        "Public resource URLs",
    ])

    arrow((430, 270), (615, 270), "HTTPS")
    arrow((985, 270), (1100, 270))
    arrow((720, 410), (540, 550))
    arrow((880, 410), (1060, 550))

    draw.text(
        (800, 825),
        "Files are uploaded to Zoho Matter_Documents attachments; Firebase Storage is not used.",
        font=regular,
        fill="#5F6B76",
        anchor="ms",
    )
    image.save(path, quality=95)


def git_value(*args):
    try:
        return subprocess.check_output(["git", *args], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def build_document():
    WORK.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram(ARCHITECTURE_PNG)

    doc = Document()
    configure_styles(doc)
    bullet_num_id, number_num_id = add_numbering(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("PLY LEGAL  |  CLIENT PORTAL DOCUMENTATION")
    set_run_font(hr, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("Technical reference  •  23 July 2026")
    set_run_font(fr, size=8.5, color=GRAY)
    fp.add_run("\t")
    page_label = fp.add_run("Page ")
    set_run_font(page_label, size=8.5, color=GRAY)
    add_page_field(fp)

    doc.core_properties.title = "Ply Legal Client Portal - Application and Technical Documentation"
    doc.core_properties.subject = "Features, technology stack, APIs, data storage, architecture, and operations"
    doc.core_properties.author = "Ply Legal Engineering"
    doc.core_properties.keywords = "Ply Legal, ValidifyPro, Next.js, Firebase, Zoho CRM, visa portal"

    commit = git_value("rev-parse", "--short", "HEAD")
    commit_date = git_value("show", "-s", "--format=%cs", "HEAD")

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = ROOT / "src" / "assets" / "Ply_Logo_black.png"
    logo_shape = p.add_run().add_picture(str(logo), width=Inches(2.7))
    set_picture_alt(logo_shape, "Ply Legal logo")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    kr = kicker.add_run("APPLICATION & TECHNICAL REFERENCE")
    set_run_font(kr, size=10.5, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Ply Legal Client Portal")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Features, architecture, APIs, data storage, configuration, and codebase guide")

    doc.add_paragraph().paragraph_format.space_after = Pt(56)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(5)
    set_run_font(meta.add_run("Product UI: Ply Legal | Client Portal"), size=10.5, color=BLACK, bold=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(5)
    set_run_font(meta2.add_run("Repository: validifypro-visa-portal"), size=10, color=GRAY)
    meta3 = doc.add_paragraph()
    meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta3.paragraph_format.space_after = Pt(5)
    set_run_font(meta3.add_run(f"Code snapshot: {commit} ({commit_date})"), size=10, color=GRAY)
    meta4 = doc.add_paragraph()
    meta4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta4.add_run("Reviewed: 23 July 2026  •  Current database mode: Firebase"), size=10, color=GRAY)

    add_section_break(doc)

    # Contents
    doc.add_heading("Contents", level=1)
    contents = [
        "1. Executive overview",
        "2. Users and product scope",
        "3. Feature catalogue",
        "4. Supported visa and intake flows",
        "5. System architecture and data flow",
        "6. Technology stack",
        "7. External APIs and integrations",
        "8. Internal API route catalogue",
        "9. Where information is saved",
        "10. Folder and code structure",
        "11. Configuration and environment variables",
        "12. Local development, build, and deployment",
        "13. Testing and verification",
        "14. Security and operational notes",
        "15. Maintenance guide and appendices",
    ]
    for item in contents:
        add_bullet(doc, bullet_num_id, item)
    add_callout(
        doc,
        "Documentation basis",
        "This guide was produced from the working repository and configuration files, not from the historical planning documents. "
        "Where a dependency or route exists but is not active, it is explicitly labelled as legacy, optional, or a stub.",
        PALE_GREEN,
        GREEN,
    )

    doc.add_heading("Document conventions", level=2)
    add_bullet(doc, bullet_num_id, "“Firestore” means the default Cloud Firestore database in region asia-east2.")
    add_bullet(doc, bullet_num_id, "“Matter” means a Zoho CRM Deal linked to a portal application through zohoId.")
    add_bullet(doc, bullet_num_id, "Paths such as applications/{appId}/data/questionnaire are Firestore document paths.")
    add_bullet(doc, bullet_num_id, "“Route-level auth” describes checks implemented by the Next.js route handler itself.")

    add_section_break(doc)

    # 1
    doc.add_heading("1. Executive overview", level=1)
    doc.add_paragraph(
        "The Ply Legal Client Portal is a responsive immigration case-management portal. It lets clients sign in, maintain a profile, "
        "view visa matters sourced from Zoho CRM, complete long multi-step questionnaires, upload requested documents, review resources, "
        "submit corrections, and exchange case messages. Staff can view client conversations through a small admin surface or a Zoho CRM widget."
    )
    add_callout(
        doc,
        "Primary design",
        "Next.js provides the web interface and server routes. Firebase provides identity and the portal’s working data. "
        "Zoho CRM remains the business system of record for contacts, matters, dependant records, requested documents, uploaded files, corrections, and CRM references.",
    )
    add_table(
        doc,
        ["Area", "Current implementation"],
        [
            ("Frontend", "Next.js App Router with React; 131 page routes"),
            ("Server API", "34 Next.js route-handler files under app/api"),
            ("Identity", "Firebase Authentication, email/password, persistent browser sessions"),
            ("Portal database", "Cloud Firestore, default database in asia-east2"),
            ("CRM", "Zoho CRM REST APIs; Australian data centre by default"),
            ("Uploads", "Zoho Matter_Documents records and their Attachments"),
            ("Questionnaires", "110 intake pages across partner, protection, and temporary-work flows"),
            ("Deployment", "Replit autoscale configuration; Next.js listens on port 5000"),
        ],
        [2200, 7160],
        font_size=9.2,
        first_col_bold=True,
    )
    architecture_shape = doc.add_picture(str(ARCHITECTURE_PNG), width=Inches(6.5))
    set_picture_alt(
        architecture_shape,
        "Architecture diagram showing users connecting to the Next.js portal, which integrates with Firebase, Zoho CRM, and external helper services.",
    )
    caption = doc.add_paragraph(style="Small Note")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run("Figure 1. Current runtime architecture and data ownership.")

    # 2
    doc.add_heading("2. Users and product scope", level=1)
    doc.add_heading("2.1 Client applicant", level=2)
    for text in [
        "Signs in only when the email exists in Zoho Contacts and Portal_Access is Active.",
        "Completes profile, address, dependant, and visa-questionnaire information.",
        "Views all linked visa applications and their completion progress.",
        "Uploads evidence against Zoho Matter Document requests.",
        "Reads shared or visa-specific resources, submits document corrections, and exchanges messages.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    doc.add_heading("2.2 Portal administrator or case staff", level=2)
    for text in [
        "Uses the admin conversation list and conversation detail pages when the Firestore user role is admin.",
        "Can reply from the portal or from the Zoho CRM widget through X-Admin-Key protected endpoints.",
        "Receives client message references and unread flags in the Zoho Client_Messages module.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    doc.add_heading("2.3 Zoho workflow integration", level=2)
    for text in [
        "Can provision or disable Firebase users through webhook-style endpoints.",
        "Pushes contact updates into Firestore and receives profile, dependant, questionnaire-status, document, correction, and message updates.",
        "Uses a separate access-token broker URL; OAuth refresh credentials are not stored in this codebase.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    doc.add_heading("2.4 Public review recipient", level=2)
    doc.add_paragraph(
        "A standalone route at /review-pdf/{matterId} renders questionnaire responses and supports browser print-to-PDF. "
        "It does not require a Firebase session; Firestore access therefore depends on the application’s publicReviewAccess flag."
    )

    # 3
    doc.add_heading("3. Feature catalogue", level=1)
    feature_rows = [
        ("Access and sessions", "Zoho eligibility gate, Firebase email/password login, browser-local persistence, sign-out, access-denied handling."),
        ("Password lifecycle", "Forced first-login password change, forgot-password token generation, reset webhook, token hashing, Firebase token revocation."),
        ("Profile management", "Personal details, address, contact data, dependants, validation, fetch from Zoho, and two-way Zoho sync."),
        ("Application dashboard", "Loads Firestore applications, refreshes linked Zoho Deals, maps visa types/statuses, removes duplicates/stale records."),
        ("Questionnaire workspace", "Progress dashboard, ordered intake navigation, save draft, continue, validation, page completion, and percentage tracking."),
        ("Applicant modelling", "Main applicant, spouse/partner, children, migrating dependants, non-migrating family, sponsor, and per-profile form sections."),
        ("Cross-application reuse", "Imports matched applicant and shared questionnaire sections; includes a 482-to-186 import helper."),
        ("Dynamic definitions", "Code-based questionnaire definition with optional active Firestore override; currently used for the 482 character page."),
        ("Document upload", "Reads Matter_Documents, validates type/size, creates or matches request records, uploads attachments, marks Awaiting Approval."),
        ("Resources", "Authenticated shared resources plus visa-template notes and external links grouped by category."),
        ("Document review", "Displays a document-review pane; creates and edits Zoho Corrections linked to the Matter."),
        ("Messaging", "Firebase conversation history and unread flags; Zoho Client_Messages reference record; admin portal and CRM widget access."),
        ("Review and PDF", "Searchable/collapsible questionnaire review with browser print output; separate matter-ID route."),
        ("Admin automation", "Create/re-enable users, revoke access, receive Zoho contact updates, and count questionnaires."),
    ]
    add_table(doc, ["Feature", "What it provides"], feature_rows, [2500, 6860], font_size=8.8, first_col_bold=True)

    doc.add_heading("3.1 Current application navigation", level=2)
    doc.add_paragraph(
        "The contextual application sidebar currently exposes Questionnaire, Upload Documents, Resources, and Document Review. "
        "Additional implemented pages exist for messages, tasks, review, and deliverables, but they are not listed in the current contextual sidebar."
    )
    add_callout(
        doc,
        "Deprecated message endpoints",
        "/api/messages/create, /api/messages/fetch, and /api/messages/upload-attachment return HTTP 410. "
        "The supported messaging endpoints are /api/chat/messages and /api/chat/send; message attachments are not implemented.",
        PALE_GOLD,
        GOLD,
    )

    # 4
    doc.add_heading("4. Supported visa and intake flows", level=1)
    add_table(
        doc,
        ["Public slug", "Internal flow", "Pages", "Scope"],
        [
            ("186", "temporary-work", "Shared 38-page implementation", "Employer Nomination Scheme; rewrite injects __subclass=186."),
            ("482", "temporary-work", "Shared 38-page implementation", "Skills in Demand; rewrite injects __subclass=482."),
            ("820 / partner", "partner", "43 pages", "Partner, spouse, children, sponsor, relationship, and all-applicant sections."),
            ("866 / protection", "protection", "29 pages", "Protection applicant, spouse, children, employment, relationship, and shared sections."),
        ],
        [1200, 1850, 1500, 4810],
        font_size=8.7,
    )

    doc.add_heading("4.1 Temporary-work questionnaire", level=2)
    for text in [
        "Getting started and included-applicant selection.",
        "Main applicant: details, other names, identity, employment, education, language, skills, and contact details.",
        "Spouse/partner and child sections, including identity, education, language, custody, and other details.",
        "Migrating dependants and non-migrating family members.",
        "All applicants: addresses, countries of residence, contact details, visas, travel history, health, and character.",
        "Review/submit with strict completion checks and Zoho questionnaire-status update.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    doc.add_heading("4.2 Partner questionnaire", level=2)
    for text in [
        "Main applicant, spouse/partner, children, and included dependants.",
        "Family sponsor: details, identity, family, circumstances, addresses, contacts, previous sponsorship, travel, and character.",
        "Relationship evidence: current relationship, relationship details, previous relationships, and supporting witnesses.",
        "All applicants: addresses, contacts, visas, travel, future travel/address, health, character, and personal contacts.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    doc.add_heading("4.3 Protection questionnaire", level=2)
    for text in [
        "Main applicant, spouse/partner, children, current relationship, and employment-offer/business details.",
        "All applicants: addresses, contacts, visas, travel history, future travel, future addresses, health, character, and contacts.",
        "Review/submit uses the same shared draft and completion model as the other flows.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    # 5
    doc.add_heading("5. System architecture and data flow", level=1)
    doc.add_heading("5.1 Login and application synchronisation", level=2)
    login_steps = [
        "The browser calls /api/auth/verify-zoho with the email address.",
        "The server checks Zoho Contacts and requires Portal_Access = Active.",
        "Firebase Authentication validates the email/password account.",
        "The portal loads users/{uid} from Firestore and refreshes the contact’s related Zoho Deals.",
        "Deals are normalised into applications/{appId}; the Zoho Deal ID is stored as zohoId.",
    ]
    for step in login_steps:
        add_numbered(doc, number_num_id, step)

    doc.add_heading("5.2 Questionnaire save and submit", level=2)
    for step in [
        "A page loads the application draft from applications/{appId}/data/questionnaire.",
        "Forms use React Hook Form and page-specific validation; selected profile sections can auto-save after a 500 ms debounce.",
        "Save operations persist first and update Valtio state after success.",
        "Completion keys and the calculated percentage are stored in applications/{appId}/data/completion.",
        "The linked Zoho Deal’s Questionnaires_Status becomes In Progress or Submitted.",
    ]:
        add_numbered(doc, number_num_id, step)

    doc.add_heading("5.3 Documents and messages", level=2)
    for text in [
        "Document requirements are read from the Zoho Deal’s Matter_Documents related list. Uploads are attached to the Matter_Documents record, not to Firebase.",
        "Messages are authoritative in Firestore conversations/{conversationId}/messages. Zoho receives a lightweight Client_Messages reference with preview, unread flags, timestamps, and portal link.",
        "Corrections are created and updated directly in the Zoho Corrections module and linked to the Matter.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    # 6
    doc.add_heading("6. Technology stack", level=1)
    active_stack = [
        ("Framework", "Next.js 16.0.10, App Router, webpack builds, React Strict Mode"),
        ("UI runtime", "React 18.3.1 and React DOM 18.3.1"),
        ("Language", "Primarily JavaScript/JSX; TypeScript 5.6 compiler with allowJs=true and strict=false"),
        ("Styling", "Tailwind CSS 3.4, PostCSS, custom Ply Legal fonts, responsive utility classes"),
        ("Component primitives", "Radix UI, shadcn-style local components, Lucide icons, class-variance-authority"),
        ("Forms and validation", "React Hook Form, Zod, @hookform/resolvers"),
        ("Client state", "Valtio stores; TanStack React Query provider and fetch helper"),
        ("Backend services", "Firebase JS SDK 12.4 and Firebase Admin SDK 13.5"),
        ("CRM integration", "Axios, form-data, and a custom ZohoCRMClient"),
        ("Identifiers and dates", "nanoid and date-fns"),
        ("Testing", "Playwright 1.60 plus small Node-style unit tests"),
        ("Package and runtime", "pnpm lockfile; application targets Node.js 20+, Firebase Functions declare Node.js 22"),
    ]
    add_table(doc, ["Layer", "Technology"], active_stack, [2300, 7060], font_size=9, first_col_bold=True)

    doc.add_heading("6.1 Installed but not active in the primary runtime", level=2)
    for text in [
        "PostgreSQL/Neon/Drizzle: dependencies and drizzle.config.ts exist, but the Postgres adapter is a TODO stub and shared/schema.ts is absent.",
        "Firebase Functions: functions/index.js contains the starter template and exports no active business function.",
        "Express, Passport, WebSocket, Vite, and Wouter packages remain installed from an earlier architecture but are not the current Next.js application path.",
        "Firebase Storage is configured, but storage.rules denies every read and write.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    # 7
    doc.add_heading("7. External APIs and integrations", level=1)
    external_rows = [
        ("Zoho CRM REST API v7", "https://www.zohoapis.{datacenter}/crm/v7", "Contacts, Deals, related lists, CRUD, search, and COQL. Default datacenter is com.au."),
        ("Zoho attachment API v8", "https://www.zohoapis.{datacenter}/crm/v8", "Uploads files to Matter_Documents/{id}/Attachments."),
        ("Zoho token broker", "ACCESSTOKEN_URL or ZOHO_ACCESS_TOKEN_URL", "Returns an OAuth access token; cached in server memory for 50 minutes."),
        ("Firebase Authentication", "Firebase JS/Admin SDK", "Email/password sign-in, account provisioning, password changes, disabling users, and token revocation."),
        ("Cloud Firestore", "Firebase JS/Admin SDK", "Portal profiles, applications, questionnaire drafts, completion, messages, resources, and definitions."),
        ("Firestore REST API", "firestore.googleapis.com/v1", "The Zoho Deal refresh route queries and writes application documents using the caller’s ID token."),
        ("Password-reset email webhook", "PORTAL_RESET_EMAIL_WEBHOOK_URL", "Receives email, resetLink, and contactId. If absent, the link is generated but not emailed."),
        ("External resources", "Stored URLs, including Zoho WorkDrive links", "The Resources page opens active external links; the portal does not copy those files."),
        ("Brand asset CDN", "cdn.prod.website-files.com", "Provides the favicon declared by the root layout."),
    ]
    add_table(doc, ["Integration", "Endpoint/configuration", "Use"], external_rows, [2200, 3000, 4160], font_size=8.3, first_col_bold=True)

    doc.add_heading("7.1 Zoho modules used", level=2)
    zoho_rows = [
        ("Contacts", "Portal eligibility, profile fields, password flags/reset token, owner/account metadata."),
        ("Deals", "Visa matters/applications, Stage, visa type, closing date, questionnaire status, documents_json."),
        ("Partner_Dependents", "Spouse, children, other dependants, and non-migrating family members."),
        ("Matter_Documents", "Requested-document records, status, comments, serial order, and file attachments."),
        ("Corrections", "Field-level review issues linked to a Matter."),
        ("Client_Messages", "Conversation reference, latest preview/time, unread flags, Firebase conversation ID, and portal link."),
    ]
    add_table(doc, ["Module", "Portal usage"], zoho_rows, [2400, 6960], font_size=9, first_col_bold=True)
    add_callout(
        doc,
        "Not used",
        "The codebase contains no payment, mapping, analytics, generative-AI, or SMS API integration.",
        PALE_GREEN,
        GREEN,
    )

    # 8
    doc.add_heading("8. Internal API route catalogue", level=1)
    doc.add_paragraph(
        "All routes are implemented as Next.js App Router route handlers. “None” below means the handler does not verify a Firebase token, "
        "admin role, or shared key itself; the absence of a route-level check should not be treated as public authorisation."
    )

    api_groups = [
        ("Authentication and administration", [
            ("POST", "/api/auth/verify-zoho", "Gate login by Zoho contact and Portal_Access.", "None"),
            ("POST", "/api/auth/forgot-password", "Create hashed reset token in Zoho and call email webhook.", "None"),
            ("POST", "/api/auth/reset-password", "Validate token, update Firebase password/profile, clear Zoho flags.", "None"),
            ("POST", "/api/auth/password-changed", "Clear temporary-password/reset flags in Zoho.", "None"),
            ("POST", "/api/admin/create-user", "Create or re-enable Firebase user and Firestore profile from Zoho.", "Bearer secret if set"),
            ("POST", "/api/admin/revoke-access", "Disable Firebase user, revoke tokens, update profile.", "Bearer secret if set"),
            ("POST", "/api/webhooks/zoho-contact-update", "Push Zoho contact/profile changes into Firestore.", "Bearer secret if set"),
        ]),
        ("Applications, profiles, and questionnaires", [
            ("POST", "/api/applications/fetch-zoho-deals", "Fetch related Deals and reconcile Firestore applications.", "ID token in body for Firestore"),
            ("GET/PATCH", "/api/deals/{dealId}", "Read documents_json or update Questionnaires_Status.", "None"),
            ("POST", "/api/profile/fetch-zoho", "Fetch Contact, dependants, and Deals; optionally save profile.", "None"),
            ("POST", "/api/profile/sync-zoho", "Create/update Contact and synchronise dependants.", "None"),
            ("GET", "/api/profile/zoho-population-status", "Diagnostic profile/Zoho population status.", "None"),
            ("GET/POST", "/api/intake/dependents", "Read Zoho dependants or save selected/excluded IDs.", "None"),
            ("POST", "/api/intake/sync-dependent", "Create, update, or delete Partner_Dependents records.", "None"),
            ("GET", "/api/questionnaires/count", "Count applications with meaningful questionnaire data.", "Firebase bearer"),
        ]),
        ("Messaging", [
            ("GET", "/api/chat/messages", "Read up to 50 messages and clear reader unread flag.", "Firebase bearer + ownership"),
            ("POST", "/api/chat/send", "Write message/conversation and update Zoho reference.", "Firebase bearer + ownership"),
            ("GET", "/api/chat/admin/conversations", "List latest 100 conversations; optional unread filter.", "Firebase bearer + admin role"),
            ("GET", "/api/chat/widget/messages", "Read up to 100 messages by Zoho Deal ID.", "X-Admin-Key"),
            ("POST", "/api/chat/widget/send", "Write CRM-widget admin reply and update Zoho reference.", "X-Admin-Key"),
            ("POST", "/api/messages/create", "Deprecated; returns HTTP 410.", "None"),
            ("GET", "/api/messages/fetch", "Deprecated; returns HTTP 410.", "None"),
            ("POST", "/api/messages/upload-attachment", "Deprecated; returns HTTP 410.", "None"),
        ]),
        ("Documents, corrections, resources, and review", [
            ("GET", "/api/uploads/matter-documents", "List/sort Deal-related Matter_Documents.", "None"),
            ("POST", "/api/uploads/zoho", "Validate and upload <=5 MB file to Zoho attachment.", "None"),
            ("POST", "/api/uploads/matter-documents/comment", "Update Matter_Documents Comments.", "None"),
            ("GET/POST/PATCH", "/api/corrections", "List, create, or edit Matter-linked Zoho Corrections.", "None"),
            ("GET", "/api/resources/shared", "Return active shared resources for an authorised application.", "Firebase bearer + ownership"),
            ("GET", "/api/resources/template", "Return active visa-template notes and external links.", "Firebase bearer + ownership"),
            ("GET", "/api/review-pdf/application/{matterId}", "Find a Firestore application by Zoho Matter ID.", "None; Firestore rules apply"),
            ("GET", "/api/review-pdf/application/{matterId}/draft", "Return questionnaire draft for public review.", "None; Firestore rules apply"),
        ]),
        ("Diagnostics", [
            ("GET", "/api/test/zoho-get", "Read a Zoho Contact for integration testing.", "None"),
            ("POST", "/api/test/zoho-update", "Update a Zoho Contact for integration testing.", "None"),
            ("GET", "/api/test/zoho-verify-token", "Fetch token and verify Zoho Contacts module access.", "None"),
        ]),
    ]
    for group_name, rows in api_groups:
        doc.add_heading(group_name, level=2)
        add_table(doc, ["Method", "Route", "Purpose", "Auth"], rows, [900, 3100, 3860, 1500], font_size=7.8)

    # 9
    doc.add_heading("9. Where information is saved", level=1)
    storage_rows = [
        ("User credentials", "Firebase Authentication", "Passwords are managed by Firebase; the portal stores no plaintext password."),
        ("Portal user profile", "Firestore users/{uid}", "Personal/contact data, role, access flags, Zoho IDs, sync metadata, dependant summary."),
        ("Visa application shell", "Firestore applications/{appId}", "Owner, type/code, status, reference, timestamps, zohoId, public review flag."),
        ("Questionnaire answers", "Firestore applications/{appId}/data/questionnaire", "Nested answers, profiles, per-profile sections, selected dependants, visa context."),
        ("Completion/progress", "Firestore applications/{appId}/data/completion", "Page completion booleans and completionPercentage."),
        ("Preferences", "Firestore users/{uid}/preferences/settings", "Current prefill setting."),
        ("Legacy current draft", "Firestore users/{uid}/drafts/current", "Only used by a real-time subscription path; application-scoped draft is primary."),
        ("Zoho contact snapshot", "Firestore users/{uid}/zoho/contact", "Server-synchronised Zoho data and timestamps."),
        ("Application messages", "Firestore conversations/{applicationId}/messages/{messageId}", "Message body, sender role/UID, and created time."),
        ("Conversation summary", "Firestore conversations/{applicationId}", "Client/matter references, preview, latest time, and unread flags."),
        ("Shared resources", "Firestore resources/{resourceId}", "Active global notes and external links."),
        ("Visa resource templates", "Firestore resourceTemplates/{slug}/items/{itemId}", "Active categorized notes and external links."),
        ("Dynamic questionnaire definitions", "Firestore questionnaireDefinitions/{id}/pages/{pageId}", "Optional active override for code-defined pages."),
        ("Uploaded client files", "Zoho Matter_Documents/{id}/Attachments", "PDF/JPG/PNG/DOC/DOCX/TXT, maximum 5 MB in the upload handler."),
        ("Matter and CRM records", "Zoho Contacts/Deals/Partner_Dependents/Corrections/Client_Messages", "CRM system-of-record entities and references."),
        ("Browser session", "Firebase Auth IndexedDB persistence", "Keeps the authenticated session across browser/tab closes."),
        ("Fallback browser data", "localStorage adapter", "Used only when configured or when adapter initialisation falls back."),
    ]
    add_table(doc, ["Information", "Primary location", "Details"], storage_rows, [2300, 3000, 4060], font_size=7.9, first_col_bold=True)

    doc.add_heading("9.1 Firestore path map", level=2)
    add_code_block(doc, """users/{uid}
├── profile and access metadata
├── preferences/settings
├── drafts/current                         (legacy/current-draft listener)
└── zoho/contact                           (Zoho snapshot)

applications/{appId}
├── application metadata and zohoId
└── data/
    ├── questionnaire                      (primary draft and applicant answers)
    ├── completion                         (page flags and percentage)
    └── {dataType}                         (adapter-supported application data)

conversations/{applicationId}
└── messages/{messageId}

resources/{resourceId}
resourceTemplates/{visaSlug}/items/{itemId}
questionnaireDefinitions/{definitionId}/pages/{pageId}""")

    doc.add_heading("9.2 Browser localStorage fallback keys", level=2)
    add_code_block(doc, """ply_session
ply_user
ply_user_profile
intake_draft
intake_prefill
ply:applications
application-scoped draft, completion, and data keys generated by the adapter
ply-theme                                  (ThemeProvider)""")
    add_callout(
        doc,
        "Important",
        "Firebase Storage is not a document repository in this application. storage.rules denies all reads and writes. "
        "Changing upload storage requires a deliberate feature and security design, not only a Firebase configuration change.",
        PALE_GOLD,
        GOLD,
    )

    # 10
    doc.add_heading("10. Folder and code structure", level=1)
    add_code_block(doc, """validifypro-visa-portal/
├── app/
│   ├── api/                     34 route handlers
│   ├── applications/            Application workspace
│   ├── intake/
│   │   ├── partner/             43 pages
│   │   ├── protection/          29 pages
│   │   └── temporary-work/      38 pages
│   ├── admin/
│   ├── profile/
│   └── review-pdf/
├── src/
│   ├── components/
│   ├── hooks/
│   ├── lib/
│   ├── stores/
│   ├── reuseable/
│   └── assets/
├── public/
├── e2e/
├── scripts/
├── functions/
├── docs/
├── knowledgebase/
├── attached_assets/
├── firebase.json
├── firestore.rules
├── storage.rules
├── next.config.js
├── package.json
└── .replit""")

    doc.add_heading("10.1 Key implementation files", level=2)
    key_files = [
        ("app/layout.js", "Root metadata and provider shell."),
        ("app/providers.jsx", "React Query, theme, AuthGuard, and navigation-loading providers."),
        ("next.config.js", "Visa-route rewrites, widget CORS headers, build aliases, and environment exposure."),
        ("src/lib/firebase.js", "Browser Firebase app, Auth, Firestore, and persistent login."),
        ("src/lib/firebase-admin.js", "Server Firebase Admin initialisation."),
        ("src/lib/serverAuth.js", "Firebase token verification, role checks, and widget admin-key verification."),
        ("src/lib/zohoClient.js", "Zoho token cache and CRM CRUD/related-record/attachment helpers."),
        ("src/lib/adapters/firebase.js", "Current persistence adapter and application/draft/profile operations."),
        ("src/stores/draftStore.js", "Questionnaire state, save/completion/import, applicant management, and CRM sync."),
        ("src/lib/routes.js", "Visa-specific navigation, rewrites, and completion paths."),
        ("src/lib/questionnaires/", "Definition-based questionnaire loading, validation, and Firestore override."),
        ("app/api/", "Server routes integrating Firebase, Zoho, resources, messaging, and uploads."),
        ("firestore.rules", "Owner-based client access and public-review exception."),
        ("e2e/", "Playwright end-to-end questionnaire coverage."),
    ]
    add_table(doc, ["File/folder", "Responsibility"], key_files, [3300, 6060], font_size=8.6, first_col_bold=True)

    doc.add_heading("10.2 Route aliases", level=2)
    rewrite_rows = [
        ("/applications/186/{appId}/intake/*", "/intake/temporary-work/*?applicationId={appId}&__subclass=186"),
        ("/applications/482/{appId}/intake/*", "/intake/temporary-work/*?applicationId={appId}&__subclass=482"),
        ("/applications/820/{appId}/intake/*", "/intake/partner/*?applicationId={appId}"),
        ("/applications/866/{appId}/intake/*", "/intake/protection/*?applicationId={appId}"),
        ("/applications/partner/{appId}/intake/*", "/intake/partner/*?applicationId={appId}"),
        ("/applications/protection/{appId}/intake/*", "/intake/protection/*?applicationId={appId}"),
    ]
    add_table(doc, ["Public application route", "Internal destination"], rewrite_rows, [4300, 5060], font_size=8.2)

    # 11
    doc.add_heading("11. Configuration and environment variables", level=1)
    env_rows = [
        ("NEXT_PUBLIC_DATABASE_TYPE", "Required", "Set to firebase in the current environment. Other accepted values: localStorage, postgres."),
        ("NEXT_PUBLIC_FIREBASE_API_KEY", "Required", "Firebase web configuration."),
        ("NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN", "Required", "Firebase web configuration."),
        ("NEXT_PUBLIC_FIREBASE_PROJECT_ID", "Required", "Firebase web/Admin project and Firestore REST URL."),
        ("NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET", "Required by initialiser", "Configured even though Storage access is deny-all."),
        ("NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID", "Required", "Firebase web configuration."),
        ("NEXT_PUBLIC_FIREBASE_APP_ID", "Required", "Firebase web configuration."),
        ("FIREBASE_SERVICE_ACCOUNT_KEY", "Required for server routes", "JSON service account supplied as a secret; used by Firebase Admin."),
        ("ACCESSTOKEN_URL", "Zoho required", "Preferred Zoho access-token broker URL."),
        ("ZOHO_ACCESS_TOKEN_URL", "Zoho fallback", "Backward-compatible token broker variable."),
        ("ZOHO_DATACENTER", "Optional", "Defaults to com.au."),
        ("PORTAL_ADMIN_KEY", "Widget required", "X-Admin-Key expected from the Zoho CRM widget."),
        ("ZOHO_WEBHOOK_SECRET", "Production required", "Bearer secret for user provisioning, revocation, and contact-update webhook."),
        ("PORTAL_RESET_EMAIL_WEBHOOK_URL", "Password reset", "Webhook that sends the generated reset link."),
        ("NEXT_PUBLIC_APP_URL / APP_URL", "Optional", "Base URL used to build password-reset links; localhost fallback exists."),
        ("NEXT_PUBLIC_BASE_URL", "Optional", "Base URL included in Zoho admin chat links."),
        ("DATABASE_URL", "Inactive adapter", "Required only by Drizzle tooling; Postgres adapter is not implemented."),
        ("NEXT_DIST_DIR", "Optional", "Overrides the Next.js build directory."),
        ("E2E_PORT / CI", "Test only", "Playwright port and CI behaviour."),
    ]
    add_table(doc, ["Variable", "Use", "Notes"], env_rows, [3100, 1800, 4460], font_size=7.7, first_col_bold=True)

    doc.add_heading("11.1 Configuration files", level=2)
    config_rows = [
        ("next.config.js", "App rewrites, widget CORS headers, aliases, body-size limit, and public Firebase variables."),
        ("firebase.json", "Default Firestore database in asia-east2, Storage rules, and functions codebase."),
        ("firestore.rules", "Owner access, active-resource reads, admin definition/template writes, and explicit public-review reads."),
        ("storage.rules", "Deny all."),
        ("tailwind.config.js/.ts", "Theme, fonts, colours, and component scanning."),
        ("playwright.config.js", "E2E server and browser configuration."),
        ("drizzle.config.ts", "Dormant PostgreSQL migration configuration."),
        (".replit", "Node 20 module, dev/start commands, port 5000, and autoscale deployment."),
    ]
    add_table(doc, ["File", "Purpose"], config_rows, [2700, 6660], font_size=8.8, first_col_bold=True)

    # 12
    doc.add_heading("12. Local development, build, and deployment", level=1)
    doc.add_heading("12.1 Prerequisites", level=2)
    for text in [
        "Node.js 20 or newer for the Next.js application; Node.js 22 for Firebase Functions if that package is deployed.",
        "pnpm is preferred because pnpm-lock.yaml is committed.",
        "A Firebase project with Email/Password authentication enabled, Firestore created, and the documented environment variables.",
        "A Zoho token-broker endpoint and CRM modules/fields matching the integration.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    doc.add_heading("12.2 Common commands", level=2)
    add_code_block(doc, """pnpm install
pnpm dev                    # Next.js dev server on port 5000
pnpm build                  # production webpack build
pnpm start                  # production server on port 5000
pnpm check                  # TypeScript no-emit check
pnpm test:e2e               # Playwright tests
pnpm questionnaires:seed    # seed active questionnaire definitions into Firestore
pnpm db:push                # Drizzle push; only relevant if Postgres work is completed""")

    doc.add_heading("12.3 Deployment", level=2)
    doc.add_paragraph(
        "The committed .replit file configures an autoscale deployment that runs npm run build and npm run start, with application port 5000 exposed externally. "
        "A Netlify Next.js plugin is installed as a development dependency, but no Netlify configuration was found. Firebase Hosting is not configured."
    )
    add_callout(
        doc,
        "Firebase Functions",
        "The functions package targets Node.js 22 and caps each function at 10 instances, but functions/index.js currently exports no deployed business handler.",
        PALE_GREEN,
        GREEN,
    )

    # 13
    doc.add_heading("13. Testing and verification", level=1)
    test_rows = [
        ("e2e/questionnaire-186.spec.js", "End-to-end Employer Nomination questionnaire flow."),
        ("e2e/questionnaire-482.spec.js", "End-to-end Skills in Demand questionnaire flow."),
        ("e2e/support/questionnaireHarness.js", "Shared Playwright setup and questionnaire helpers."),
        ("src/lib/adapters/__tests__/factory.test.js", "Adapter selection."),
        ("src/lib/adapters/__tests__/localStorage.test.js", "Browser fallback persistence."),
        ("src/lib/questionnaires/__tests__/validation.node.test.js", "Questionnaire-definition validation."),
        ("src/stores/__tests__/draftStore.nonMigrating.test.js", "Non-migrating member draft behaviour."),
    ]
    add_table(doc, ["Test", "Coverage"], test_rows, [4300, 5060], font_size=8.7, first_col_bold=True)
    doc.add_paragraph(
        "Only Playwright and the TypeScript checker have package scripts. The small Node tests need an explicit runner/command if they are to be part of CI. "
        "Current E2E coverage is strongest for subclasses 186 and 482; partner, protection, messaging, uploads, webhooks, and access control need dedicated automated coverage."
    )

    # 14
    doc.add_heading("14. Security and operational notes", level=1)
    add_callout(
        doc,
        "Immediate action",
        "The repository currently tracks .env and two Firebase service-account JSON files even though .gitignore now excludes .env. "
        "Treat the credentials as exposed: rotate them, remove the files from Git history, and move all secrets to the deployment secret store.",
        PALE_RED,
        RED,
    )
    security_rows = [
        ("Critical", "Tracked secrets", ".env, firebase-service-account.json, and validify-pro-test-firebase-adminsdk-fbsvc-4f4b7a7e17.json are tracked.", "Rotate credentials and purge Git history."),
        ("High", "Missing route-level auth", "Several profile, dependant, deal, correction, upload, review, and test routes do not verify caller identity/ownership.", "Add central auth/authorisation; disable diagnostic routes in production."),
        ("High", "Webhook fail-open", "Webhook/admin routes enforce ZOHO_WEBHOOK_SECRET only when the variable is configured.", "Fail closed when the secret is absent."),
        ("High", "Credential logging", "firebase-admin.js logs the length and first 50 characters of FIREBASE_SERVICE_ACCOUNT_KEY.", "Remove all credential-derived logs."),
        ("High", "Upload route", "/api/uploads/zoho validates file type/size but does not authenticate or verify Deal ownership.", "Require Firebase token and application ownership."),
        ("Medium", "Public review", "Matter-ID review routes have no route auth; Firestore publicReviewAccess is the effective gate.", "Keep the flag false by default and add signed/time-limited access."),
        ("Medium", "Debug ingestion", "draftStore and one partner page call a hard-coded 127.0.0.1:7242 ingest endpoint.", "Remove debug calls before release."),
        ("Medium", "CORS scope", "Zoho widget routes return Access-Control-Allow-Origin: *.", "Restrict allowed origins if the widget host is stable."),
        ("Medium", "Reset token lifetime", "The Zoho reset-token hash has no expiry timestamp in the current flow.", "Store and enforce an expiry and single-use audit data."),
        ("Low", "Dormant code", "Postgres, Express/Vite remnants, deprecated routes, backups, and test pages expand the maintenance surface.", "Remove once confirmed unused."),
    ]
    add_table(doc, ["Severity", "Area", "Current observation", "Recommended action"], security_rows, [1050, 1700, 3530, 3080], font_size=7.3)

    doc.add_heading("14.1 Existing safeguards", level=2)
    for text in [
        "Firestore owner rules protect users and application subcollections; unrecognised paths are denied.",
        "Authenticated chat routes verify Firebase ID tokens and application ownership; admin conversation listing checks the admin role.",
        "Zoho widget routes require PORTAL_ADMIN_KEY.",
        "Login fails closed if Zoho eligibility cannot be verified.",
        "Password-reset responses avoid disclosing whether an email exists.",
        "Uploads enforce a 5 MB size limit and an allow-list of document MIME types.",
        "Firebase Storage denies all reads and writes.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    # 15
    doc.add_heading("15. Maintenance guide and appendices", level=1)
    doc.add_heading("15.1 Where to make common changes", level=2)
    maintenance_rows = [
        ("Add/edit a static intake page", "app/intake/{flow}/.../page.js plus src/lib/routes.js."),
        ("Change questionnaire route order", "src/lib/routes.js and completion/validation helpers."),
        ("Add a definition-driven question", "src/lib/questionnaires, questionnaire renderer components, then run questionnaires:seed."),
        ("Change data persistence", "src/lib/adapters/firebase.js; keep Firestore rules and server routes aligned."),
        ("Change profile/CRM mapping", "src/lib/zohoClient.js and app/api/profile routes."),
        ("Change Deal/application mapping", "app/api/applications/fetch-zoho-deals/route.js and src/lib/visaDisplay.js."),
        ("Change document uploads", "application uploads page, /api/uploads routes, and Zoho Matter_Documents fields."),
        ("Change messaging", "application/admin message pages, /api/chat routes, and Zoho Client_Messages fields."),
        ("Change resources", "/api/resources routes, Firestore resource collections, and resource page."),
        ("Change access control", "src/lib/serverAuth.js, each route handler, firestore.rules, and AuthGuard."),
        ("Change deployment", ".replit, package scripts, runtime secrets, and Firebase/Zoho configuration."),
    ]
    add_table(doc, ["Task", "Primary locations"], maintenance_rows, [3100, 6260], font_size=8.3, first_col_bold=True)

    doc.add_heading("15.2 Route surface summary", level=2)
    summary_rows = [
        ("Top-level page routes", "131", "Authentication, profile, application workspace, intake, admin, and review."),
        ("Intake pages", "110", "Partner 43; protection 29; temporary-work 38."),
        ("API route files", "34", "Auth/admin, applications, profile, intake, chat, resources, uploads, review, and tests."),
        ("Application workspace pages", "9 child pages", "Questionnaire, uploads, resources, corrections, messages, tasks, review, deliverables, plus layout."),
        ("Admin pages", "2", "Conversation list and detail."),
        ("Public review pages", "1", "Matter-ID questionnaire review and print."),
    ]
    add_table(doc, ["Surface", "Count", "Notes"], summary_rows, [2800, 1300, 5260], font_size=8.8, first_col_bold=True)

    doc.add_heading("15.3 Repository snapshot notes", level=2)
    for text in [
        f"Documented commit: {commit}, dated {commit_date}.",
        "The working tree was inspected without reading or reproducing secret values.",
        "Historical markdown plans and attached screenshots remain useful background, but the code and active configuration were treated as authoritative.",
        "The product metadata names the application “Ply Legal | Client Portal”; the repository/workspace name remains validifypro-visa-portal.",
    ]:
        add_bullet(doc, bullet_num_id, text)

    add_callout(
        doc,
        "End of guide",
        "Update this document after material changes to authentication, persistence, Zoho modules, route access control, supported visa flows, or deployment.",
        PALE_GREEN,
        GREEN,
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
